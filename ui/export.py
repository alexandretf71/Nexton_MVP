from io import BytesIO
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_BRAND_BLUE = RGBColor(0x33, 0x7B, 0xFF)
_LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "Nexton Logo_black.png")
_TAGLINE = "Specially developed, quickly and intelligently, by Alexandre T. F."


def _add_document_header(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    try:
        run_logo = para.add_run()
        run_logo.add_picture(_LOGO_PATH, height=Inches(0.3))
    except Exception:
        run_name = para.add_run("nexton")
        run_name.bold = True
        run_name.font.size = Pt(13)
        run_name.font.name = "K2D"
        run_name.font.color.rgb = _BRAND_BLUE

    para.add_run("    ")

    run_tag = para.add_run(_TAGLINE)
    run_tag.italic = True
    run_tag.font.size = Pt(8)
    run_tag.font.color.rgb = None  # inherit theme colour (grey in most themes)

    # Thin bottom border on header paragraph via OOXML
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_heading(doc: Document, number: int, title: str) -> None:
    doc.add_paragraph(f"{number}. {title}", style="Heading 1")


def _shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
        _shade_cell(hdr_cells[i], "D9D9D9")
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val) if val is not None else ""
    doc.add_paragraph()


def _add_org_chart_bullets(doc: Document, roles: list[dict]) -> None:
    # reports_to values without a matching role act as roots (e.g. "AI Division Director").
    role_names = {r.get("role") for r in roles}
    children: dict[str, list[dict]] = {}
    for r in roles:
        children.setdefault(r.get("reports_to", ""), []).append(r)

    bullet_styles = ["List Bullet", "List Bullet 2", "List Bullet 3"]
    visited: set[str] = set()

    def _walk(name: str, depth: int) -> None:
        if name in visited:
            return
        visited.add(name)
        for child in children.get(name, []):
            style = bullet_styles[min(depth, len(bullet_styles) - 1)]
            doc.add_paragraph(
                f"{child.get('role')} (×{child.get('count')}, {child.get('seniority')})",
                style=style,
            )
            _walk(child.get("role", ""), depth + 1)

    for root in sorted(set(children) - role_names):
        doc.add_paragraph(root, style="List Bullet")
        _walk(root, 1)


def generate_blueprint_docx(bp: dict) -> bytes:
    doc = Document()
    _add_document_header(doc)

    # ── Cover block ───────────────────────────────────────────────────────────
    doc.add_paragraph("AI Implementation Blueprint", style="Title")
    solution_type = bp.get("recommended_solution_type", "")
    if solution_type:
        doc.add_paragraph(solution_type, style="Subtitle")
    generated_at = (bp.get("generated_at") or "")[:19].replace("T", " ")
    provider = bp.get("provider_used", "")
    doc.add_paragraph(f"Generated: {generated_at} UTC   |   Provider: {provider}")
    doc.add_paragraph()

    # ── 1. Business Problem Summary ───────────────────────────────────────────
    _add_section_heading(doc, 1, "Business Problem Summary")
    doc.add_paragraph(bp.get("business_problem_summary") or "", style="Normal")

    # ── 2. Operational Pain Points ────────────────────────────────────────────
    _add_section_heading(doc, 2, "Operational Pain Points")
    for point in bp.get("operational_pain_points") or []:
        doc.add_paragraph(str(point), style="List Bullet")

    # ── 3. AI Opportunity Classification ─────────────────────────────────────
    _add_section_heading(doc, 3, "AI Opportunity Classification")
    for opp in bp.get("ai_opportunity_classification") or []:
        doc.add_paragraph(opp.get("title", ""), style="Heading 2")
        doc.add_paragraph(
            f"Impact: {opp.get('impact', '')}   |   Feasibility: {opp.get('feasibility', '')}",
            style="Normal",
        )
        doc.add_paragraph(opp.get("description", ""), style="Normal")

    # ── 4. Recommended Solution Type ──────────────────────────────────────────
    _add_section_heading(doc, 4, "Recommended Solution Type")
    doc.add_paragraph(bp.get("recommended_solution_type") or "", style="Normal")

    # ── 5. Suggested Architecture ─────────────────────────────────────────────
    _add_section_heading(doc, 5, "Suggested Architecture")
    for line in (bp.get("suggested_architecture") or "").split("\n"):
        doc.add_paragraph(line, style="Normal")

    # ── 6. Agentic Workflow ───────────────────────────────────────────────────
    _add_section_heading(doc, 6, "Agentic Workflow")
    for step in bp.get("agentic_workflow") or []:
        doc.add_paragraph(
            f"Step {step.get('step_number', '')}: {step.get('agent_name', '')}",
            style="Heading 2",
        )
        doc.add_paragraph(f"Action: {step.get('action', '')}", style="Normal")
        doc.add_paragraph(f"Input: {step.get('input', '')}", style="Normal")
        doc.add_paragraph(f"Output: {step.get('output', '')}", style="Normal")
        tools = ", ".join(step.get("tools_used") or [])
        doc.add_paragraph(f"Tools: {tools}", style="Normal")

    # ── 7. Data Requirements ──────────────────────────────────────────────────
    _add_section_heading(doc, 7, "Data Requirements")
    _add_table(
        doc,
        ["Source", "Availability", "Description", "Notes"],
        [
            [r.get("source"), r.get("availability"), r.get("description"), r.get("notes")]
            for r in bp.get("data_requirements") or []
        ],
    )

    # ── 8. Integration Points ─────────────────────────────────────────────────
    _add_section_heading(doc, 8, "Integration Points")
    _add_table(
        doc,
        ["System", "Type", "Complexity", "Description"],
        [
            [r.get("system"), r.get("integration_type"), r.get("complexity"), r.get("description")]
            for r in bp.get("integration_points") or []
        ],
    )

    # ── 9. Engineering Backlog ────────────────────────────────────────────────
    _add_section_heading(doc, 9, "Engineering Backlog")
    _add_table(
        doc,
        ["ID", "Title", "Priority", "Effort", "Phase", "Description"],
        [
            [r.get("id"), r.get("title"), r.get("priority"), r.get("effort"), r.get("phase"), r.get("description")]
            for r in bp.get("engineering_backlog") or []
        ],
    )

    # ── 10. Risks & Assumptions ───────────────────────────────────────────────
    _add_section_heading(doc, 10, "Risks & Assumptions")
    _add_table(
        doc,
        ["ID", "Category", "Likelihood", "Impact", "Description", "Mitigation"],
        [
            [r.get("id"), r.get("category"), r.get("likelihood"), r.get("impact"), r.get("description"), r.get("mitigation")]
            for r in bp.get("risks_and_assumptions") or []
        ],
    )

    # ── 11. Testing Plan ──────────────────────────────────────────────────────
    _add_section_heading(doc, 11, "Testing Plan")
    _add_table(
        doc,
        ["ID", "Type", "Description", "Expected Outcome"],
        [
            [r.get("id"), r.get("type"), r.get("description"), r.get("expected_outcome")]
            for r in bp.get("testing_plan") or []
        ],
    )

    # ── 12. Executive Summary ─────────────────────────────────────────────────
    _add_section_heading(doc, 12, "Executive Summary")
    doc.add_paragraph(bp.get("executive_summary") or "", style="Normal")

    # ── 13. Delivery Status Report ────────────────────────────────────────────
    _add_section_heading(doc, 13, "Delivery Status Report")
    for milestone in bp.get("delivery_status_report") or []:
        doc.add_paragraph(
            f"Week {milestone.get('week', '')}: {milestone.get('title', '')}   —   Owner: {milestone.get('owner', '')}",
            style="Heading 2",
        )
        for deliverable in milestone.get("deliverables") or []:
            doc.add_paragraph(str(deliverable), style="List Bullet")

    # ── 14. Delivery Team & Fulfillment Plan ──────────────────────────────────
    _add_section_heading(doc, 14, "Delivery Team & Fulfillment Plan")
    disclaimer = doc.add_paragraph()
    run_disc = disclaimer.add_run(
        "Note: team sizing assumes a deliberately reduced team optimized by "
        "AI-assisted coding tools (e.g. Claude Code, GitHub Copilot, Cursor). "
        "Traditional staffing for the same scope would be substantially larger."
    )
    run_disc.italic = True
    run_disc.font.size = Pt(9)
    team = bp.get("delivery_team") or {}
    roles = team.get("roles") or []
    _add_table(
        doc,
        ["Role", "Count", "Seniority", "Allocation", "Reports To", "Responsibilities"],
        [
            [r.get("role"), r.get("count"), r.get("seniority"), r.get("allocation"),
             r.get("reports_to"), r.get("responsibilities")]
            for r in roles
        ],
    )

    doc.add_paragraph("Org Chart", style="Heading 2")
    _add_org_chart_bullets(doc, roles)

    doc.add_paragraph("Fulfillment Plan", style="Heading 2")
    _add_table(
        doc,
        ["Week", "Role", "Action", "Notes"],
        [
            [fa.get("week"), fa.get("role"), fa.get("action"), fa.get("notes")]
            for fa in team.get("fulfillment_plan") or []
        ],
    )

    # ── Serialise to bytes ────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
